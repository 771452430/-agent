"""知识树 + 文档解析 + scoped Hybrid RAG 检索。

这一版不再把知识库看成扁平文档列表，而是显式引入“知识树”：
- 节点描述知识范围；
- 文档归属到某个节点；
- chunk 带上 tree metadata，检索时可按 scope 过滤。

在检索实现上，这一版采用 Hybrid RAG：
- lexical 检索负责稳定命中文件名、路径、关键词；
- vector 检索负责补充“更像语义”的相似片段；
- 最后再用 RRF（Reciprocal Rank Fusion）融合排序。

这样前端的“检索模式”和“我的 Agent”就能共享同一套知识范围能力。
"""

from __future__ import annotations

import json
import re
import shutil
import sqlite3
import subprocess
import quopri
from dataclasses import dataclass, field
from datetime import datetime, timezone
from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path, PurePosixPath
from tempfile import NamedTemporaryFile
from typing import Any
from uuid import uuid4

from ..schemas import (
    Citation,
    DirectoryUploadItem,
    KnowledgeDocument,
    KnowledgeDeleteResponse,
    KnowledgeTreeNode,
    KnowledgeTreeNodeDetail,
    KnowledgeTreeResponse,
    RetrievalCandidateDebug,
    RetrievalProfile,
    ScopeType,
)
from ..settings import AppSettings
from .embedding_service import EmbeddingService
from .provider_store import ProviderStore
from .rag_embedding_settings_service import RAGEmbeddingSettingsService
from .vector_store import ChunkVectorRecord, KnowledgeVectorStore, VectorSearchHit


ROOT_NODE_ID = "root"
ROOT_NODE_NAME = "全部知识"
SUPPORTED_DOCUMENT_TYPES = {"txt", "md", "pdf", "doc", "docx", "xlsx"}


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


@dataclass(frozen=True)
class ScoredCitation:
    """Hybrid 检索融合后的中间结果。"""

    citation: Citation
    fused_score: float
    lexical_score: float = 0.0
    vector_score: float = 0.0
    lexical_rank: int | None = None
    vector_rank: int | None = None


@dataclass(frozen=True)
class StructuredChunk:
    """结构化切片结果。

    `heading_path` 表示文档内部的小节路径，
    后续会用于 lexical 加权、证据卡片展示和 rerank 理由解释。
    """

    content: str
    heading_path: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


class _HtmlTextExtractor(HTMLParser):
    BLOCK_TAGS = {
        "p",
        "div",
        "br",
        "tr",
        "table",
        "section",
        "article",
        "ul",
        "ol",
        "li",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "blockquote",
        "td",
        "th",
    }
    SKIP_PAIRED_TAGS = {"script", "style", "head", "title", "xml"}
    SKIP_VOID_TAGS = {"meta", "link"}

    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        lowered = tag.lower()
        if lowered in self.SKIP_VOID_TAGS:
            return
        if lowered in self.SKIP_PAIRED_TAGS:
            self.skip_depth += 1
            return
        if self.skip_depth == 0 and lowered in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        lowered = tag.lower()
        if lowered in self.SKIP_PAIRED_TAGS and self.skip_depth > 0:
            self.skip_depth -= 1
            return
        if self.skip_depth == 0 and lowered in self.BLOCK_TAGS:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self.skip_depth == 0:
            self.parts.append(data)

    def get_text(self) -> str:
        return "".join(self.parts)


class KnowledgeStore:
    """统一管理知识树、文档入库和检索。"""

    def __init__(
        self,
        sqlite_path: Path,
        chroma_dir: Path,
        *,
        provider_store: ProviderStore | None = None,
        settings: AppSettings | None = None,
        rag_embedding_settings_service: RAGEmbeddingSettingsService | None = None,
    ) -> None:
        # 初始化顺序很重要：
        # 1. 先准备数据库表和旧数据兼容；
        # 2. 再修复历史文档格式；
        # 3. 最后尝试接上向量库。
        # 这样即使向量库不可用，最基础的 lexical 检索仍然可以工作。
        self.sqlite_path = sqlite_path
        self.chroma_dir = chroma_dir
        self.embedding_service = EmbeddingService(
            provider_store=provider_store,
            settings=settings,
            rag_embedding_settings_service=rag_embedding_settings_service,
        )
        self.vector_store: KnowledgeVectorStore | None = None
        self._embedding_backend = self.embedding_service.describe_runtime_selection().preferred_backend
        self._retrieval_backend = "lexical"
        self._init_db()
        self._embedding_backend = self._load_indexed_embedding_backend() or self._embedding_backend
        self._repair_legacy_doc_documents()
        self._init_vector_store()

    def _init_vector_store(self) -> None:
        """初始化 Chroma，并把已有 chunk 尽量同步进去。"""

        try:
            # 向量库初始化成功后，系统进入 hybrid 检索模式：
            # lexical 负责关键词命中，vector 负责语义补充。
            self.vector_store = KnowledgeVectorStore(self.chroma_dir)
            self._retrieval_backend = "hybrid"
            self._sync_vector_index()
        except Exception:
            # 这里故意降级而不是直接抛错，
            # 因为学习模式下“能继续用 lexical 检索”比“整个系统启动失败”更重要。
            self.vector_store = None
            self._retrieval_backend = "lexical"

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.sqlite_path)
        conn.row_factory = sqlite3.Row
        return conn

    def _table_columns(self, conn: sqlite3.Connection, table_name: str) -> set[str]:
        rows = conn.execute(f"PRAGMA table_info({table_name})").fetchall()
        return {row["name"] for row in rows}

    def _ensure_column(self, conn: sqlite3.Connection, table_name: str, column_name: str, column_sql: str) -> None:
        if column_name not in self._table_columns(conn, table_name):
            conn.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_sql}")

    def _init_db(self) -> None:
        with self._connect() as conn:
            # 这里既负责首启建表，也承担“旧版本升级补列”的责任。
            # 因为项目是教学 demo，直接在启动时做轻量迁移，阅读和维护都更直观。
            conn.executescript(
                """
                CREATE TABLE IF NOT EXISTS knowledge_tree_nodes (
                    id TEXT PRIMARY KEY,
                    name TEXT NOT NULL,
                    parent_id TEXT,
                    path TEXT NOT NULL UNIQUE,
                    created_at TEXT NOT NULL
                );

                CREATE TABLE IF NOT EXISTS knowledge_documents (
                    id TEXT PRIMARY KEY,
                    node_id TEXT NOT NULL DEFAULT 'root',
                    name TEXT NOT NULL,
                    type TEXT NOT NULL,
                    relative_path TEXT NOT NULL DEFAULT '',
                    status TEXT NOT NULL,
                    chunk_count INTEGER NOT NULL DEFAULT 0,
                    created_at TEXT NOT NULL,
                    error_message TEXT,
                    external_url TEXT,
                    metadata_json TEXT NOT NULL DEFAULT '{}'
                );

                CREATE TABLE IF NOT EXISTS knowledge_chunks (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    node_id TEXT NOT NULL DEFAULT 'root',
                    document_name TEXT NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    content TEXT NOT NULL,
                    tree_path TEXT NOT NULL DEFAULT '/',
                    relative_path TEXT NOT NULL DEFAULT '',
                    source_type TEXT NOT NULL DEFAULT 'txt',
                    heading_path TEXT NOT NULL DEFAULT '',
                    metadata_json TEXT NOT NULL DEFAULT '{}',
                    created_at TEXT NOT NULL
                );

                CREATE TABLE IF NOT EXISTS knowledge_runtime_state (
                    singleton_id INTEGER PRIMARY KEY CHECK (singleton_id = 1),
                    indexed_embedding_backend TEXT NOT NULL DEFAULT '',
                    updated_at TEXT NOT NULL
                );
                """
            )

            self._ensure_column(conn, "knowledge_documents", "node_id", "TEXT NOT NULL DEFAULT 'root'")
            self._ensure_column(conn, "knowledge_documents", "relative_path", "TEXT NOT NULL DEFAULT ''")
            self._ensure_column(conn, "knowledge_documents", "external_url", "TEXT")
            self._ensure_column(conn, "knowledge_documents", "metadata_json", "TEXT NOT NULL DEFAULT '{}'")
            self._ensure_column(conn, "knowledge_chunks", "node_id", "TEXT NOT NULL DEFAULT 'root'")
            self._ensure_column(conn, "knowledge_chunks", "tree_path", "TEXT NOT NULL DEFAULT '/'")
            self._ensure_column(conn, "knowledge_chunks", "relative_path", "TEXT NOT NULL DEFAULT ''")
            self._ensure_column(conn, "knowledge_chunks", "source_type", "TEXT NOT NULL DEFAULT 'txt'")
            self._ensure_column(conn, "knowledge_chunks", "heading_path", "TEXT NOT NULL DEFAULT ''")
            self._ensure_column(conn, "knowledge_chunks", "metadata_json", "TEXT NOT NULL DEFAULT '{}'")

            now = _utc_now().isoformat()
            conn.execute(
                """
                INSERT OR IGNORE INTO knowledge_tree_nodes (id, name, parent_id, path, created_at)
                VALUES (?, ?, NULL, '/', ?)
                """,
                (ROOT_NODE_ID, ROOT_NODE_NAME, now),
            )

            conn.execute("UPDATE knowledge_documents SET node_id = 'root' WHERE node_id IS NULL OR node_id = ''")
            conn.execute(
                "UPDATE knowledge_documents SET relative_path = name WHERE relative_path IS NULL OR relative_path = ''"
            )
            conn.execute(
                "UPDATE knowledge_documents SET metadata_json = '{}' "
                "WHERE metadata_json IS NULL OR metadata_json = ''"
            )
            conn.execute("UPDATE knowledge_chunks SET node_id = 'root' WHERE node_id IS NULL OR node_id = ''")
            conn.execute("UPDATE knowledge_chunks SET tree_path = '/' WHERE tree_path IS NULL OR tree_path = ''")
            conn.execute(
                "UPDATE knowledge_chunks SET heading_path = '' "
                "WHERE heading_path IS NULL"
            )
            conn.execute(
                "UPDATE knowledge_chunks SET metadata_json = '{}' "
                "WHERE metadata_json IS NULL OR metadata_json = ''"
            )
            conn.execute(
                """
                UPDATE knowledge_chunks
                SET relative_path = COALESCE(
                    (SELECT kd.relative_path FROM knowledge_documents kd WHERE kd.id = knowledge_chunks.document_id),
                    relative_path,
                    ''
                )
                """
            )
            conn.execute(
                """
                UPDATE knowledge_chunks
                SET source_type = COALESCE(
                    (SELECT kd.type FROM knowledge_documents kd WHERE kd.id = knowledge_chunks.document_id),
                    source_type,
                    'txt'
                )
                """
            )

    def _load_indexed_embedding_backend(self) -> str:
        with self._connect() as conn:
            row = conn.execute(
                "SELECT indexed_embedding_backend FROM knowledge_runtime_state WHERE singleton_id = 1"
            ).fetchone()
        return str(row["indexed_embedding_backend"] or "").strip() if row is not None else ""

    def _save_indexed_embedding_backend(self, backend_name: str) -> None:
        normalized = str(backend_name or "").strip()
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO knowledge_runtime_state (singleton_id, indexed_embedding_backend, updated_at)
                VALUES (1, ?, ?)
                ON CONFLICT(singleton_id) DO UPDATE SET
                    indexed_embedding_backend = excluded.indexed_embedding_backend,
                    updated_at = excluded.updated_at
                """,
                (normalized, _utc_now().isoformat()),
            )

    def _load_ready_chunk_records(self, document_id: str | None = None) -> list[ChunkVectorRecord]:
        sql = """
            SELECT kc.id, kc.document_id, kc.node_id, kc.document_name,
                   kc.content, kc.tree_path, kc.relative_path, kc.source_type,
                   kc.heading_path, kc.metadata_json
            FROM knowledge_chunks kc
            JOIN knowledge_documents kd ON kd.id = kc.document_id
            WHERE kd.status = 'ready'
        """
        params: list[str] = []
        if document_id is not None:
            sql += " AND kc.document_id = ?"
            params.append(document_id)
        sql += " ORDER BY kc.created_at ASC, kc.chunk_index ASC"
        with self._connect() as conn:
            rows = conn.execute(sql, params).fetchall()
        return [self._row_to_chunk_record(row) for row in rows]

    def _index_chunk_records(self, records: list[ChunkVectorRecord]) -> None:
        if self.vector_store is None or not records:
            return
        batch_size = 500
        last_backend = self.embedding_service.describe_runtime_selection().preferred_backend
        for start in range(0, len(records), batch_size):
            batch_records = records[start : start + batch_size]
            embeddings = self.embedding_service.embed_documents([record.content for record in batch_records])
            self.vector_store.upsert_chunks(batch_records, embeddings)
            last_backend = self.embedding_service.model_name
        self._embedding_backend = last_backend
        self._save_indexed_embedding_backend(last_backend)

    def _reconstruct_chunk_text(self, chunks: list[str], max_overlap: int = 200) -> str:
        if not chunks:
            return ""
        restored = chunks[0]
        for chunk in chunks[1:]:
            best_overlap = 0
            upper = min(max_overlap, len(restored), len(chunk))
            for size in range(upper, 19, -1):
                if restored.endswith(chunk[:size]):
                    best_overlap = size
                    break
            restored += chunk[best_overlap:]
        return restored

    def _looks_like_legacy_doc_markup(self, text: str) -> bool:
        lowered = text.lower()
        return any(
            marker in lowered
            for marker in [
                "exported from confluence",
                "content-transfer-encoding: quoted-printable",
                "content-type: multipart/related",
                "<html",
            ]
        )

    def _extract_html_segment(self, text: str) -> str:
        lowered = text.lower()
        start = lowered.find("<html")
        end = lowered.rfind("</html>")
        if start >= 0 and end > start:
            return text[start : end + len("</html>")]
        return text

    def _html_to_text(self, html_text: str) -> str:
        extractor = _HtmlTextExtractor()
        extractor.feed(html_text)
        extractor.close()
        return extractor.get_text()

    def _normalize_legacy_doc_text(self, raw_text: str) -> str:
        normalized = raw_text.replace("\r\n", "\n").replace("\r", "\n")
        if self._looks_like_legacy_doc_markup(normalized):
            decoded = quopri.decodestring(normalized.encode("utf-8", errors="ignore")).decode(
                "utf-8", errors="ignore"
            )
            if decoded.strip():
                normalized = decoded
            normalized = self._html_to_text(self._extract_html_segment(normalized))
        normalized = unescape(normalized)
        cleaned_lines: list[str] = []
        for line in normalized.splitlines():
            compact = " ".join(line.split())
            if not compact:
                continue
            lowered = compact.lower()
            if lowered.startswith(
                (
                    "date:",
                    "message-id:",
                    "mime-version:",
                    "content-type:",
                    "content-transfer-encoding:",
                    "content-location:",
                )
            ):
                continue
            if compact.startswith("------=_Part_"):
                continue
            if re.fullmatch(r"[A-Za-z0-9+/=]{80,}", compact):
                continue
            cleaned_lines.append(compact)
        return "\n".join(cleaned_lines)

    def _repair_legacy_doc_documents(self) -> None:
        with self._connect() as conn:
            doc_rows = conn.execute(
                """
                SELECT kd.*, kc.content AS first_chunk
                FROM knowledge_documents kd
                LEFT JOIN knowledge_chunks kc
                  ON kc.document_id = kd.id AND kc.chunk_index = 0
                WHERE kd.type = 'doc' AND kd.status = 'ready'
                ORDER BY kd.created_at ASC
                """
            ).fetchall()

        for row in doc_rows:
            first_chunk = str(row["first_chunk"] or "")
            if not self._looks_like_legacy_doc_markup(first_chunk):
                continue
            with self._connect() as conn:
                chunk_rows = conn.execute(
                    "SELECT content FROM knowledge_chunks WHERE document_id = ? ORDER BY chunk_index ASC",
                    (row["id"],),
                ).fetchall()
            raw_text = self._reconstruct_chunk_text([str(chunk["content"] or "") for chunk in chunk_rows])
            cleaned_text = self._normalize_legacy_doc_text(raw_text)
            if not cleaned_text or cleaned_text == raw_text:
                continue
            chunks = self._split_text(cleaned_text)
            target_node = self._require_node_row(row["node_id"])
            with self._connect() as conn:
                conn.execute("DELETE FROM knowledge_chunks WHERE document_id = ?", (row["id"],))
            self._store_chunks(
                document_id=str(row["id"]),
                node_id=str(row["node_id"]),
                document_name=str(row["name"]),
                relative_path=str(row["relative_path"] or row["name"]),
                tree_path=str(target_node["path"]),
                source_type=str(row["type"]),
                chunks=chunks,
            )
            with self._connect() as conn:
                conn.execute(
                    """
                    UPDATE knowledge_documents
                    SET chunk_count = ?, error_message = NULL
                    WHERE id = ?
                    """,
                    (len(chunks), row["id"]),
                )

    def _sync_vector_index(self) -> None:
        """把已有 ready chunk 同步到向量库。

        这个 demo 优先选择更容易讲清楚的实现：
        只要发现 SQLite 里的 chunk 数和 Chroma 不一致，就做一次全量重建。
        """

        if self.vector_store is None:
            return
        records = self._load_ready_chunk_records()
        indexed_backend = self._load_indexed_embedding_backend()
        if self.vector_store.count() == len(records) and indexed_backend != "":
            self._embedding_backend = indexed_backend
            return
        self.rebuild_vector_index()

    def rebuild_vector_index(self) -> None:
        """按当前 embedding 运行时配置全量重建知识库向量索引。"""

        if self.vector_store is None:
            return
        records = self._load_ready_chunk_records()
        self.vector_store.reset()
        if not records:
            backend = self.embedding_service.describe_runtime_selection().preferred_backend
            self._embedding_backend = backend
            self._save_indexed_embedding_backend(backend)
            return
        self._index_chunk_records(records)

    def _normalize_relative_path(self, relative_path: str | None, fallback_name: str) -> str:
        raw = (relative_path or fallback_name).replace("\\", "/").strip("/")
        return str(PurePosixPath(raw or fallback_name))

    def _build_node_path(self, parent_path: str, node_name: str) -> str:
        if parent_path == "/":
            return f"/{node_name}"
        return f"{parent_path}/{node_name}"

    def _get_node_row(self, node_id: str) -> sqlite3.Row | None:
        with self._connect() as conn:
            return conn.execute("SELECT * FROM knowledge_tree_nodes WHERE id = ?", (node_id,)).fetchone()

    def _require_node_row(self, node_id: str | None) -> sqlite3.Row:
        target = node_id or ROOT_NODE_ID
        row = self._get_node_row(target)
        if row is None:
            raise ValueError("知识树节点不存在")
        return row

    def _get_or_create_child_node(self, parent_id: str, child_name: str) -> sqlite3.Row:
        parent = self._require_node_row(parent_id)
        with self._connect() as conn:
            existing = conn.execute(
                """
                SELECT * FROM knowledge_tree_nodes
                WHERE parent_id = ? AND name = ?
                """,
                (parent["id"], child_name),
            ).fetchone()
            if existing:
                return existing

            node_id = str(uuid4())
            node_path = self._build_node_path(parent["path"], child_name)
            created_at = _utc_now().isoformat()
            conn.execute(
                """
                INSERT INTO knowledge_tree_nodes (id, name, parent_id, path, created_at)
                VALUES (?, ?, ?, ?, ?)
                """,
                (node_id, child_name, parent["id"], node_path, created_at),
            )
            row = conn.execute("SELECT * FROM knowledge_tree_nodes WHERE id = ?", (node_id,)).fetchone()
        assert row is not None
        return row

    def _row_to_document(self, row: sqlite3.Row) -> KnowledgeDocument:
        metadata = self._parse_metadata_json(row["metadata_json"] if "metadata_json" in row.keys() else "{}")
        return KnowledgeDocument(
            id=row["id"],
            node_id=row["node_id"],
            name=row["name"],
            type=row["type"],
            relative_path=row["relative_path"] or row["name"],
            status=row["status"],
            chunk_count=row["chunk_count"],
            created_at=datetime.fromisoformat(row["created_at"]),
            error_message=row["error_message"],
            external_url=row["external_url"],
            metadata=metadata,
        )

    def _parse_metadata_json(self, raw: str | None) -> dict[str, Any]:
        if raw is None or raw == "":
            return {}
        try:
            data = json.loads(raw)
        except Exception:
            return {}
        return data if isinstance(data, dict) else {}

    def _row_to_citation(self, row: sqlite3.Row) -> Citation:
        metadata = self._parse_metadata_json(row["metadata_json"] if "metadata_json" in row.keys() else "{}")
        return Citation(
            document_id=row["document_id"],
            document_name=row["document_name"],
            chunk_id=row["id"],
            snippet=(row["content"] or "")[:260],
            tree_id=row["node_id"],
            tree_path=row["tree_path"],
            relative_path=row["relative_path"],
            source_type=row["source_type"],
            heading_path=row["heading_path"] if "heading_path" in row.keys() else "",
            metadata=metadata,
        )

    def _vector_hit_to_citation(self, hit: VectorSearchHit) -> Citation:
        metadata = self._parse_metadata_json(hit.metadata_json)
        return Citation(
            document_id=hit.document_id,
            document_name=hit.document_name,
            chunk_id=hit.chunk_id,
            snippet=hit.content[:260],
            tree_id=hit.node_id,
            tree_path=hit.tree_path,
            relative_path=hit.relative_path,
            source_type=hit.source_type,
            heading_path=hit.heading_path,
            metadata=metadata,
        )

    def _row_to_chunk_record(self, row: sqlite3.Row) -> ChunkVectorRecord:
        return ChunkVectorRecord(
            chunk_id=row["id"],
            document_id=row["document_id"],
            node_id=row["node_id"],
            document_name=row["document_name"],
            content=row["content"],
            tree_path=row["tree_path"],
            relative_path=row["relative_path"],
            source_type=row["source_type"],
            heading_path=row["heading_path"] if "heading_path" in row.keys() else "",
            metadata_json=row["metadata_json"] if "metadata_json" in row.keys() else "{}",
        )

    def _serialize_tree(self) -> KnowledgeTreeResponse:
        with self._connect() as conn:
            node_rows = conn.execute("SELECT * FROM knowledge_tree_nodes ORDER BY path ASC").fetchall()
            child_count_rows = conn.execute(
                """
                SELECT parent_id, COUNT(*) AS count
                FROM knowledge_tree_nodes
                WHERE parent_id IS NOT NULL
                GROUP BY parent_id
                """
            ).fetchall()
            doc_count_rows = conn.execute(
                """
                SELECT node_id, COUNT(*) AS count
                FROM knowledge_documents
                GROUP BY node_id
                """
            ).fetchall()

        child_counts = {row["parent_id"]: row["count"] for row in child_count_rows}
        doc_counts = {row["node_id"]: row["count"] for row in doc_count_rows}
        nodes = {
            row["id"]: KnowledgeTreeNode(
                id=row["id"],
                name=row["name"],
                parent_id=row["parent_id"],
                path=row["path"],
                children_count=child_counts.get(row["id"], 0),
                document_count=doc_counts.get(row["id"], 0),
                children=[],
            )
            for row in node_rows
        }

        for node in nodes.values():
            if node.parent_id and node.parent_id in nodes:
                nodes[node.parent_id].children.append(node)

        def sort_children(current: KnowledgeTreeNode) -> None:
            current.children.sort(key=lambda item: (item.path.count("/"), item.name))
            for child in current.children:
                sort_children(child)

        root = nodes.get(ROOT_NODE_ID)
        if root is None:
            root = KnowledgeTreeNode(
                id=ROOT_NODE_ID,
                name=ROOT_NODE_NAME,
                parent_id=None,
                path="/",
                children_count=0,
                document_count=0,
                children=[],
            )
        sort_children(root)
        return KnowledgeTreeResponse(root=root)

    def get_tree(self) -> KnowledgeTreeResponse:
        return self._serialize_tree()

    def list_documents(self, node_id: str | None = None) -> list[KnowledgeDocument]:
        sql = "SELECT * FROM knowledge_documents"
        params: list[str] = []
        if node_id:
            sql += " WHERE node_id = ?"
            params.append(node_id)
        sql += " ORDER BY created_at DESC"
        with self._connect() as conn:
            rows = conn.execute(sql, params).fetchall()
        return [self._row_to_document(row) for row in rows]

    def get_document(self, document_id: str) -> KnowledgeDocument | None:
        with self._connect() as conn:
            row = conn.execute("SELECT * FROM knowledge_documents WHERE id = ?", (document_id,)).fetchone()
        return self._row_to_document(row) if row is not None else None

    def update_document_metadata(self, document_id: str, *, external_url: str | None) -> KnowledgeDocument:
        current = self.get_document(document_id)
        if current is None:
            raise ValueError("知识文件不存在")
        normalized_external_url = (external_url or "").strip() or None
        with self._connect() as conn:
            conn.execute(
                """
                UPDATE knowledge_documents
                SET external_url = ?
                WHERE id = ?
                """,
                (normalized_external_url, document_id),
            )
        updated = self.get_document(document_id)
        assert updated is not None
        return updated

    def get_document_external_urls(self, document_ids: list[str]) -> dict[str, str]:
        normalized_ids = [document_id for document_id in document_ids if document_id.strip() != ""]
        if len(normalized_ids) == 0:
            return {}
        placeholders = ",".join("?" for _ in normalized_ids)
        with self._connect() as conn:
            rows = conn.execute(
                f"SELECT id, external_url FROM knowledge_documents WHERE id IN ({placeholders})",
                normalized_ids,
            ).fetchall()
        return {
            str(row["id"]): str(row["external_url"]).strip()
            for row in rows
            if row["external_url"] is not None and str(row["external_url"]).strip() != ""
        }

    def is_supported_document(self, filename: str) -> bool:
        file_type = Path(filename).suffix.lower().lstrip(".")
        return file_type in SUPPORTED_DOCUMENT_TYPES

    def find_document_by_relative_path(self, node_id: str, relative_path: str) -> KnowledgeDocument | None:
        fallback_name = PurePosixPath(relative_path).name or "untitled.txt"
        normalized_relative_path = self._normalize_relative_path(relative_path, fallback_name)
        with self._connect() as conn:
            row = conn.execute(
                """
                SELECT * FROM knowledge_documents
                WHERE node_id = ? AND relative_path = ?
                ORDER BY created_at DESC
                LIMIT 1
                """,
                (node_id, normalized_relative_path),
            ).fetchone()
        return self._row_to_document(row) if row is not None else None

    def upsert_document(
        self,
        *,
        parent_node_id: str | None,
        relative_path: str,
        file_bytes: bytes,
        external_url: str | None = None,
    ) -> tuple[KnowledgeDocument, bool]:
        fallback_name = PurePosixPath(relative_path).name or "untitled.txt"
        normalized_relative_path = self._normalize_relative_path(relative_path, fallback_name)
        path_parts = list(PurePosixPath(normalized_relative_path).parts)
        if len(path_parts) == 0:
            raise ValueError("relative_path 不能为空")

        current_node = self._require_node_row(parent_node_id)
        for folder_name in path_parts[:-1]:
            current_node = self._get_or_create_child_node(current_node["id"], folder_name)

        file_name = path_parts[-1]
        existing = self.find_document_by_relative_path(current_node["id"], normalized_relative_path)
        updated = existing is not None
        if existing is not None:
            self.delete_document(existing.id)

        document = self.ingest_document(
            filename=file_name,
            file_bytes=file_bytes,
            node_id=current_node["id"],
            relative_path=normalized_relative_path,
        )
        if external_url is not None:
            document = self.update_document_metadata(document.id, external_url=external_url)
        return document, updated

    def create_node(self, name: str, parent_id: str | None = None) -> KnowledgeTreeNode:
        clean_name = name.strip()
        if not clean_name:
            raise ValueError("节点名称不能为空")
        row = self._get_or_create_child_node(parent_id or ROOT_NODE_ID, clean_name)
        tree = self._serialize_tree()

        def find_node(current: KnowledgeTreeNode) -> KnowledgeTreeNode | None:
            if current.id == row["id"]:
                return current
            for child in current.children:
                match = find_node(child)
                if match:
                    return match
            return None

        node = find_node(tree.root)
        assert node is not None
        return node

    def get_node_detail(self, node_id: str) -> KnowledgeTreeNodeDetail:
        tree = self._serialize_tree()
        target: KnowledgeTreeNode | None = None

        def walk(current: KnowledgeTreeNode) -> None:
            nonlocal target
            if current.id == node_id:
                target = current
                return
            for child in current.children:
                walk(child)

        walk(tree.root)
        if target is None:
            raise ValueError("知识树节点不存在")

        scope_node_ids = self._resolve_scope_node_ids("tree_recursive", node_id)
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT COUNT(*) AS count
                FROM knowledge_documents
                WHERE node_id IN (%s)
                """
                % ",".join("?" for _ in scope_node_ids),
                list(scope_node_ids),
            ).fetchone()
        recursive_document_count = int(rows["count"]) if rows else 0
        return KnowledgeTreeNodeDetail(
            node=target,
            children=target.children,
            documents=self.list_documents(node_id),
            recursive_document_count=recursive_document_count,
            recursive_children_count=max(len(scope_node_ids) - 1, 0),
        )

    def _delete_chunk_vectors(self, chunk_ids: list[str]) -> None:
        if self.vector_store is None or not chunk_ids:
            return
        try:
            self.vector_store.delete_chunks(chunk_ids)
        except Exception:
            self._sync_vector_index()

    def delete_document(self, document_id: str) -> KnowledgeDeleteResponse:
        with self._connect() as conn:
            document_row = conn.execute("SELECT * FROM knowledge_documents WHERE id = ?", (document_id,)).fetchone()
            if document_row is None:
                raise ValueError("知识文件不存在")
            chunk_rows = conn.execute(
                "SELECT id FROM knowledge_chunks WHERE document_id = ? ORDER BY chunk_index ASC",
                (document_id,),
            ).fetchall()
            chunk_ids = [str(row["id"]) for row in chunk_rows]
            deleted_chunk_count = conn.execute(
                "DELETE FROM knowledge_chunks WHERE document_id = ?",
                (document_id,),
            ).rowcount
            deleted_document_count = conn.execute(
                "DELETE FROM knowledge_documents WHERE id = ?",
                (document_id,),
            ).rowcount

        self._delete_chunk_vectors(chunk_ids)
        return KnowledgeDeleteResponse(
            message=f"已删除文件《{document_row['name']}》",
            deleted_node_count=0,
            deleted_document_count=int(deleted_document_count),
            deleted_chunk_count=int(deleted_chunk_count),
        )

    def delete_node(self, node_id: str) -> KnowledgeDeleteResponse:
        target = self._require_node_row(node_id)
        if target["id"] == ROOT_NODE_ID:
            raise ValueError("根节点不支持删除")

        scope_node_ids = sorted(self._resolve_scope_node_ids("tree_recursive", node_id))
        if not scope_node_ids:
            raise ValueError("知识树节点不存在")
        with self._connect() as conn:
            placeholders = ",".join("?" for _ in scope_node_ids)
            chunk_rows = conn.execute(
                f"SELECT id FROM knowledge_chunks WHERE node_id IN ({placeholders})",
                scope_node_ids,
            ).fetchall()
            chunk_ids = [str(row["id"]) for row in chunk_rows]
            deleted_chunk_count = conn.execute(
                f"DELETE FROM knowledge_chunks WHERE node_id IN ({placeholders})",
                scope_node_ids,
            ).rowcount
            deleted_document_count = conn.execute(
                f"DELETE FROM knowledge_documents WHERE node_id IN ({placeholders})",
                scope_node_ids,
            ).rowcount
            deleted_node_count = conn.execute(
                f"DELETE FROM knowledge_tree_nodes WHERE id IN ({placeholders})",
                scope_node_ids,
            ).rowcount

        self._delete_chunk_vectors(chunk_ids)
        return KnowledgeDeleteResponse(
            message=f"已删除目录《{target['name']}》及其下内容",
            deleted_node_count=int(deleted_node_count),
            deleted_document_count=int(deleted_document_count),
            deleted_chunk_count=int(deleted_chunk_count),
        )

    def has_documents(self, scope_type: ScopeType = "global", scope_id: str | None = None) -> bool:
        where_sql = "WHERE status = 'ready'"
        params: list[str] = []
        if scope_type == "tree_recursive":
            scope_node_ids = list(self._resolve_scope_node_ids(scope_type, scope_id))
            if not scope_node_ids:
                return False
            where_sql += f" AND node_id IN ({','.join('?' for _ in scope_node_ids)})"
            params.extend(scope_node_ids)
        with self._connect() as conn:
            row = conn.execute(f"SELECT COUNT(*) AS count FROM knowledge_documents {where_sql}", params).fetchone()
        return bool(row and row["count"] > 0)

    def ingest_document(
        self,
        filename: str,
        file_bytes: bytes,
        *,
        node_id: str | None = None,
        relative_path: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> KnowledgeDocument:
        """把单个文件挂到某个知识树节点下，并完成解析、切分和入库。"""
        target_node = self._require_node_row(node_id)
        document_id = str(uuid4())
        file_type = Path(filename).suffix.lower().lstrip(".") or "txt"
        created_at = _utc_now().isoformat()
        normalized_relative_path = self._normalize_relative_path(relative_path, filename)
        document_metadata = dict(metadata or {})

        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO knowledge_documents (
                    id, node_id, name, type, relative_path, status, chunk_count, created_at, error_message, metadata_json
                )
                VALUES (?, ?, ?, ?, ?, 'processing', 0, ?, NULL, ?)
                """,
                (
                    document_id,
                    target_node["id"],
                    filename,
                    file_type,
                    normalized_relative_path,
                    created_at,
                    json.dumps(document_metadata, ensure_ascii=False),
                ),
            )

        try:
            raw_text = self._extract_text(file_type=file_type, file_bytes=file_bytes)
            chunks = self._split_text(raw_text)
            chunk_records = self._store_chunks(
                document_id=document_id,
                node_id=target_node["id"],
                document_name=filename,
                relative_path=normalized_relative_path,
                tree_path=target_node["path"],
                source_type=file_type,
                chunks=chunks,
                document_metadata=document_metadata,
            )
            vector_warning: str | None = None
            try:
                self._index_chunk_records(chunk_records)
            except Exception as exc:
                vector_warning = f"向量索引暂未完成，当前先保留 lexical 检索：{exc}"
            with self._connect() as conn:
                conn.execute(
                    """
                    UPDATE knowledge_documents
                    SET status = 'ready', chunk_count = ?, error_message = ?
                    WHERE id = ?
                    """,
                    (len(chunks), vector_warning, document_id),
                )
        except Exception as exc:
            with self._connect() as conn:
                conn.execute(
                    """
                    UPDATE knowledge_documents
                    SET status = 'error', error_message = ?
                    WHERE id = ?
                    """,
                    (str(exc), document_id),
                )

        with self._connect() as conn:
            row = conn.execute("SELECT * FROM knowledge_documents WHERE id = ?", (document_id,)).fetchone()
        assert row is not None
        return self._row_to_document(row)

    def ingest_directory(self, parent_node_id: str | None, files: list[DirectoryUploadItem]) -> list[KnowledgeDocument]:
        """按 relative_path 批量导入目录。

        这一步是知识树体验的关键：前端并不是先传树结构，
        而是把目录里的文件和 relative_path 一起传进来，后端再根据路径逐级建树。
        """
        parent = self._require_node_row(parent_node_id)
        results: list[KnowledgeDocument] = []
        for item in files:
            normalized_relative_path = self._normalize_relative_path(item.relative_path, item.file_name)
            path_parts = list(PurePosixPath(normalized_relative_path).parts)
            current_node = parent
            for folder_name in path_parts[:-1]:
                current_node = self._get_or_create_child_node(current_node["id"], folder_name)
            file_name = path_parts[-1] if path_parts else item.file_name
            results.append(
                self.ingest_document(
                    file_name=file_name,
                    file_bytes=BytesIO(__import__("base64").b64decode(item.content_base64.encode("utf-8"))).getvalue(),
                    node_id=current_node["id"],
                    relative_path=normalized_relative_path,
                )
            )
        return results

    def _extract_text(self, file_type: str, file_bytes: bytes) -> str:
        if file_type in {"txt", "md"}:
            return file_bytes.decode("utf-8", errors="ignore")

        if file_type == "pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:  # pragma: no cover
                raise RuntimeError("需要安装 pypdf 才能解析 PDF") from exc
            with NamedTemporaryFile(suffix=".pdf") as temp_file:
                temp_file.write(file_bytes)
                temp_file.flush()
                reader = PdfReader(temp_file.name)
                return "\n".join(page.extract_text() or "" for page in reader.pages)

        if file_type == "docx":
            try:
                from docx import Document as DocxDocument
            except ImportError as exc:  # pragma: no cover
                raise RuntimeError("需要安装 python-docx 才能解析 DOCX") from exc
            with NamedTemporaryFile(suffix=".docx") as temp_file:
                temp_file.write(file_bytes)
                temp_file.flush()
                document = DocxDocument(temp_file.name)
                return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())

        if file_type == "doc":
            return self._extract_legacy_word_text(file_bytes)

        if file_type == "xlsx":
            try:
                from openpyxl import load_workbook
            except ImportError as exc:  # pragma: no cover
                raise RuntimeError("需要安装 openpyxl 才能解析 XLSX") from exc

            workbook = load_workbook(filename=BytesIO(file_bytes), data_only=True)
            lines: list[str] = []
            for worksheet in workbook.worksheets:
                rows = [
                    [str(cell).strip() if cell is not None else "" for cell in row]
                    for row in worksheet.iter_rows(values_only=True)
                ]
                rows = [row for row in rows if any(cell for cell in row)]
                if not rows:
                    continue
                headers = rows[0]
                lines.append(f"[Sheet] {worksheet.title}")
                if any(header for header in headers):
                    lines.append("表头：" + " | ".join(header or f"列{index + 1}" for index, header in enumerate(headers)))
                for row_index, row in enumerate(rows[1:], start=1):
                    pairs = []
                    for col_index, value in enumerate(row):
                        if not value:
                            continue
                        header = headers[col_index] if col_index < len(headers) and headers[col_index] else f"列{col_index + 1}"
                        pairs.append(f"{header}: {value}")
                    if pairs:
                        lines.append(f"第{row_index}行：{'；'.join(pairs)}")
                lines.append("")
            return "\n".join(lines)

        raise RuntimeError(f"暂不支持的文档类型：{file_type}")

    def _extract_legacy_word_text(self, file_bytes: bytes) -> str:
        textutil_path = shutil.which("textutil")
        if textutil_path is None:
            raise RuntimeError("当前环境缺少 textutil，暂时无法解析 DOC 文件")

        with NamedTemporaryFile(suffix=".doc") as temp_file:
            temp_file.write(file_bytes)
            temp_file.flush()
            result = subprocess.run(
                [textutil_path, "-convert", "txt", "-stdout", "-encoding", "UTF-8", temp_file.name],
                capture_output=True,
                text=True,
            )
        if result.returncode != 0:
            message = result.stderr.strip() or result.stdout.strip() or "未知错误"
            raise RuntimeError(f"DOC 解析失败：{message}")
        return self._normalize_legacy_doc_text(result.stdout)

    def _structured_blocks(self, text: str) -> list[StructuredChunk]:
        """把原始文本先拆成带标题路径的结构块。"""

        normalized_lines = [line.rstrip() for line in str(text or "").splitlines()]
        heading_stack: list[str] = []
        blocks: list[StructuredChunk] = []
        current_lines: list[str] = []
        current_heading_path = ""
        in_code_block = False

        def flush_current_block() -> None:
            nonlocal current_lines, current_heading_path
            content = "\n".join(line.strip() for line in current_lines if line.strip() != "").strip()
            if content == "":
                current_lines = []
                return
            blocks.append(
                StructuredChunk(
                    content=content,
                    heading_path=current_heading_path,
                    metadata={"heading_path": current_heading_path},
                )
            )
            current_lines = []

        for raw_line in normalized_lines:
            stripped = raw_line.strip()
            if stripped.startswith("```"):
                if not in_code_block:
                    flush_current_block()
                    current_heading_path = " / ".join(heading_stack)
                    current_lines = [raw_line]
                    in_code_block = True
                else:
                    current_lines.append(raw_line)
                    flush_current_block()
                    in_code_block = False
                continue

            if in_code_block:
                current_lines.append(raw_line)
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if heading_match:
                flush_current_block()
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                if title != "":
                    heading_stack = heading_stack[: level - 1]
                    heading_stack.append(title)
                current_heading_path = " / ".join(heading_stack)
                continue

            if stripped == "":
                flush_current_block()
                current_heading_path = " / ".join(heading_stack)
                continue

            if re.match(r"^([-*]|\d+\.)\s+", stripped):
                flush_current_block()
                blocks.append(
                    StructuredChunk(
                        content=stripped,
                        heading_path=" / ".join(heading_stack),
                        metadata={"heading_path": " / ".join(heading_stack), "block_type": "list_item"},
                    )
                )
                current_heading_path = " / ".join(heading_stack)
                continue

            current_heading_path = " / ".join(heading_stack)
            current_lines.append(raw_line)

        flush_current_block()
        return blocks

    def _sliding_window_text(self, text: str, *, chunk_size: int, overlap: int) -> list[str]:
        chunks: list[str] = []
        start = 0
        while start < len(text):
            end = min(len(text), start + chunk_size)
            chunks.append(text[start:end])
            if end >= len(text):
                break
            start = max(0, end - overlap)
        return chunks

    def _split_text(self, text: str, chunk_size: int = 520, overlap: int = 80) -> list[StructuredChunk]:
        """按标题、段落、列表和代码块做结构化切片。"""

        blocks = self._structured_blocks(text)
        if not blocks:
            raise RuntimeError("文档解析后没有可索引内容")

        chunks: list[StructuredChunk] = []
        current_parts: list[str] = []
        current_heading_path = ""
        for block in blocks:
            block_text = block.content.strip()
            if block_text == "":
                continue

            # 如果单个结构块过长，单独按滑窗切开，避免标题路径丢失。
            if len(block_text) > chunk_size:
                if current_parts:
                    merged = "\n\n".join(current_parts).strip()
                    chunks.append(
                        StructuredChunk(
                            content=merged,
                            heading_path=current_heading_path,
                            metadata={"heading_path": current_heading_path},
                        )
                    )
                    current_parts = []
                for window in self._sliding_window_text(block_text, chunk_size=chunk_size, overlap=overlap):
                    chunks.append(
                        StructuredChunk(
                            content=window.strip(),
                            heading_path=block.heading_path,
                            metadata=dict(block.metadata),
                        )
                    )
                current_heading_path = ""
                continue

            candidate_parts = [*current_parts, block_text]
            candidate_text = "\n\n".join(candidate_parts).strip()
            if current_parts and (
                len(candidate_text) > chunk_size or block.heading_path != current_heading_path
            ):
                merged = "\n\n".join(current_parts).strip()
                chunks.append(
                    StructuredChunk(
                        content=merged,
                        heading_path=current_heading_path,
                        metadata={"heading_path": current_heading_path},
                    )
                )
                overlap_seed = merged[-overlap:].strip()
                current_parts = [overlap_seed] if overlap_seed != "" else []

            current_heading_path = block.heading_path
            current_parts.append(block_text)

        if current_parts:
            chunks.append(
                StructuredChunk(
                    content="\n\n".join(current_parts).strip(),
                    heading_path=current_heading_path,
                    metadata={"heading_path": current_heading_path},
                )
            )
        return [chunk for chunk in chunks if chunk.content.strip() != ""]

    def _store_chunks(
        self,
        *,
        document_id: str,
        node_id: str,
        document_name: str,
        relative_path: str,
        tree_path: str,
        source_type: str,
        chunks: list[StructuredChunk],
        document_metadata: dict[str, Any] | None = None,
    ) -> list[ChunkVectorRecord]:
        now = _utc_now().isoformat()
        records: list[ChunkVectorRecord] = []
        base_metadata = dict(document_metadata or {})
        rows = [
            (
                chunk_id,
                document_id,
                node_id,
                document_name,
                index,
                chunk.content,
                tree_path,
                relative_path,
                source_type,
                chunk.heading_path,
                json.dumps(
                    {
                        **base_metadata,
                        **chunk.metadata,
                    },
                    ensure_ascii=False,
                ),
                now,
            )
            for index, chunk in enumerate(chunks)
            for chunk_id in [str(uuid4())]
        ]
        for row in rows:
            records.append(
                ChunkVectorRecord(
                    chunk_id=row[0],
                    document_id=row[1],
                    node_id=row[2],
                    document_name=row[3],
                    content=row[5],
                    tree_path=row[6],
                    relative_path=row[7],
                    source_type=row[8],
                    heading_path=row[9],
                    metadata_json=row[10],
                )
            )
        with self._connect() as conn:
            conn.executemany(
                """
                INSERT INTO knowledge_chunks (
                    id, document_id, node_id, document_name, chunk_index, content,
                    tree_path, relative_path, source_type, heading_path, metadata_json, created_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                rows,
            )
        return records

    def _resolve_scope_node_ids(self, scope_type: ScopeType, scope_id: str | None) -> set[str]:
        if scope_type in {"none", "global"}:
            return {row.id for row in []}
        start_id = scope_id or ROOT_NODE_ID
        with self._connect() as conn:
            rows = conn.execute("SELECT id, parent_id FROM knowledge_tree_nodes").fetchall()
        children_map: dict[str, list[str]] = {}
        known_ids = {row["id"] for row in rows}
        if start_id not in known_ids:
            return set()
        for row in rows:
            parent_id = row["parent_id"]
            if parent_id is not None:
                children_map.setdefault(parent_id, []).append(row["id"])
        visited: set[str] = set()
        stack = [start_id]
        while stack:
            current = stack.pop()
            if current in visited:
                continue
            visited.add(current)
            stack.extend(children_map.get(current, []))
        return visited

    def _query_tokens(self, query: str) -> list[str]:
        normalized = query.lower().strip()
        if not normalized:
            return []
        tokens = {normalized}
        tokens.update(re.findall(r"[a-z0-9_]+", normalized))
        chinese_parts = re.findall(r"[\u4e00-\u9fff]{2,}", normalized)
        tokens.update(chinese_parts)
        for part in chinese_parts:
            if len(part) > 2:
                tokens.update(part[index : index + 2] for index in range(0, len(part) - 1))
        return [token for token in tokens if token]

    def _is_identifier_token(self, token: str) -> bool:
        normalized = token.strip().lower()
        if normalized == "":
            return False
        return bool(re.search(r"\d", normalized) or re.search(r"[._:/-]", normalized) or len(normalized) >= 8)

    def _score_text_matches(self, *, text: str, tokens: list[str], exact_query: str, base_weight: float) -> float:
        lowered = text.lower()
        if lowered == "":
            return 0.0

        score = 0.0
        if exact_query != "" and exact_query in lowered:
            score += base_weight * 3.0
        for token in tokens:
            if token not in lowered:
                continue
            token_weight = base_weight * (2.0 if self._is_identifier_token(token) else 1.0)
            score += token_weight
        return score

    def _profile_priority_boost(
        self,
        *,
        retrieval_profile: RetrievalProfile,
        query: str,
        candidate: RetrievalCandidateDebug,
    ) -> float:
        if retrieval_profile != "support_issue":
            return 0.0

        metadata = candidate.metadata
        boost = 0.0
        if metadata.get("source") == "approved_case" or (candidate.relative_path or "").startswith("支持案例库/"):
            boost += 0.18

        normalized_query = query.lower().strip()
        identifier_terms = [token for token in self._query_tokens(query) if self._is_identifier_token(token)]
        title_text = " ".join(
            [
                candidate.document_name or "",
                candidate.relative_path or "",
                candidate.heading_path or "",
            ]
        ).lower()
        for token in identifier_terms:
            if token in title_text:
                boost += 0.08
        if normalized_query != "" and normalized_query in title_text:
            boost += 0.05
        return min(boost, 0.35)

    def _row_to_candidate(self, row: sqlite3.Row) -> RetrievalCandidateDebug:
        citation = self._row_to_citation(row)
        return RetrievalCandidateDebug(
            chunk_id=citation.chunk_id,
            document_id=citation.document_id,
            document_name=citation.document_name,
            snippet=citation.snippet,
            tree_id=citation.tree_id,
            tree_path=citation.tree_path,
            relative_path=citation.relative_path,
            source_type=citation.source_type,
            heading_path=citation.heading_path,
            metadata=citation.metadata,
        )

    def _vector_hit_to_candidate(self, hit: VectorSearchHit) -> RetrievalCandidateDebug:
        citation = self._vector_hit_to_citation(hit)
        return RetrievalCandidateDebug(
            chunk_id=citation.chunk_id,
            document_id=citation.document_id,
            document_name=citation.document_name,
            snippet=citation.snippet,
            tree_id=citation.tree_id,
            tree_path=citation.tree_path,
            relative_path=citation.relative_path,
            source_type=citation.source_type,
            heading_path=citation.heading_path,
            metadata=citation.metadata,
        )

    def search_candidates(
        self,
        query: str,
        *,
        limit: int = 6,
        scope_type: ScopeType = "global",
        scope_id: str | None = None,
        retrieval_profile: RetrievalProfile = "default",
    ) -> list[RetrievalCandidateDebug]:
        """返回带打分信息的候选片段，供上层做多 query 合并和 rerank。"""

        lexical_hits = self._lexical_search(
            query,
            limit=max(limit * 3, 12),
            scope_type=scope_type,
            scope_id=scope_id,
            retrieval_profile=retrieval_profile,
        )
        vector_hits = self._vector_search(
            query,
            limit=max(limit * 3, 12),
            scope_type=scope_type,
            scope_id=scope_id,
            retrieval_profile=retrieval_profile,
        )

        if lexical_hits and vector_hits:
            return self._fuse_search_results(query, lexical_hits, vector_hits, limit, retrieval_profile=retrieval_profile)
        if vector_hits:
            candidates: list[RetrievalCandidateDebug] = []
            for rank, hit in enumerate(vector_hits[:limit], start=1):
                candidate = self._vector_hit_to_candidate(hit)
                vector_score = self._vector_score_from_distance(hit.distance)
                fused_score = self._rrf_score(rank) + self._profile_priority_boost(
                    retrieval_profile=retrieval_profile,
                    query=query,
                    candidate=candidate,
                )
                candidates.append(
                    candidate.model_copy(
                        update={
                            "vector_score": vector_score,
                            "fused_score": fused_score,
                        }
                    )
                )
            return candidates
        candidates = []
        for rank, (lexical_score, row) in enumerate(lexical_hits[:limit], start=1):
            candidate = self._row_to_candidate(row)
            fused_score = self._rrf_score(rank) + self._profile_priority_boost(
                retrieval_profile=retrieval_profile,
                query=query,
                candidate=candidate,
            )
            candidates.append(
                candidate.model_copy(
                    update={
                        "lexical_score": float(lexical_score),
                        "fused_score": fused_score,
                    }
                )
            )
        return candidates

    def search(
        self,
        query: str,
        *,
        limit: int = 6,
        scope_type: ScopeType = "global",
        scope_id: str | None = None,
        retrieval_profile: RetrievalProfile = "default",
    ) -> list[Citation]:
        """执行按范围过滤后的 Hybrid 检索。

        这里最重要的学习点不是检索算法本身，而是 scope 的落点：
        先根据 scope_type / scope_id 算出允许命中的 node 集合，
        再在这些节点下分别做 lexical / vector 召回，并最终融合排序。
        """
        candidates = self.search_candidates(
            query,
            limit=limit,
            scope_type=scope_type,
            scope_id=scope_id,
            retrieval_profile=retrieval_profile,
        )
        return [
            Citation(
                document_id=item.document_id,
                document_name=item.document_name,
                chunk_id=item.chunk_id,
                snippet=item.snippet,
                tree_id=item.tree_id,
                tree_path=item.tree_path,
                relative_path=item.relative_path,
                source_type=item.source_type,
                heading_path=item.heading_path,
                metadata=item.metadata,
            )
            for item in candidates
        ]

    def _lexical_search(
        self,
        query: str,
        *,
        limit: int,
        scope_type: ScopeType,
        scope_id: str | None,
        retrieval_profile: RetrievalProfile,
    ) -> list[tuple[float, sqlite3.Row]]:
        """保留原有 lexical 检索，作为 Hybrid 的一个召回来源。"""

        sql = """
            SELECT kc.id, kc.document_id, kc.node_id, kc.document_name, kc.content,
                   kc.tree_path, kc.relative_path, kc.source_type, kc.heading_path, kc.metadata_json
            FROM knowledge_chunks kc
            JOIN knowledge_documents kd ON kd.id = kc.document_id
            WHERE kd.status = 'ready'
        """
        params: list[str] = []
        if scope_type == "tree_recursive":
            scope_node_ids = list(self._resolve_scope_node_ids(scope_type, scope_id))
            if not scope_node_ids:
                return []
            sql += f" AND kc.node_id IN ({','.join('?' for _ in scope_node_ids)})"
            params.extend(scope_node_ids)
        sql += " ORDER BY kc.created_at DESC"

        with self._connect() as conn:
            rows = conn.execute(sql, params).fetchall()

        tokens = self._query_tokens(query)
        exact_query = query.strip().lower()
        scored: list[tuple[float, sqlite3.Row]] = []
        for row in rows:
            title_text = " ".join(
                [row["document_name"] or "", row["heading_path"] or ""]
            ).lower()
            path_text = " ".join([row["relative_path"] or "", row["tree_path"] or ""]).lower()
            body_text = str(row["content"] or "").lower()
            score = 0.0
            score += self._score_text_matches(text=title_text, tokens=tokens, exact_query=exact_query, base_weight=2.6)
            score += self._score_text_matches(text=path_text, tokens=tokens, exact_query=exact_query, base_weight=1.8)
            score += self._score_text_matches(text=body_text, tokens=tokens, exact_query=exact_query, base_weight=0.9)
            if score > 0:
                candidate = self._row_to_candidate(row)
                score += self._profile_priority_boost(
                    retrieval_profile=retrieval_profile,
                    query=query,
                    candidate=candidate,
                )
                scored.append((score, row))

        scored.sort(key=lambda item: item[0], reverse=True)
        return scored[:limit]

    def _vector_search(
        self,
        query: str,
        *,
        limit: int,
        scope_type: ScopeType,
        scope_id: str | None,
        retrieval_profile: RetrievalProfile,
    ) -> list[VectorSearchHit]:
        """向量召回：先做 scope 过滤，再做 Chroma 相似度查询。"""

        if self.vector_store is None or not query.strip():
            return []
        scope_node_ids: list[str] | None = None
        if scope_type == "tree_recursive":
            scope_node_ids = list(self._resolve_scope_node_ids(scope_type, scope_id))
            if not scope_node_ids:
                return []
        try:
            query_embedding = self.embedding_service.embed_query(query)
            indexed_backend = self._load_indexed_embedding_backend()
            query_backend = self.embedding_service.model_name
            if indexed_backend != "" and query_backend != indexed_backend:
                # 查询侧如果因为 provider 不可用回退成 hashing，就不要拿它去查
                # 另一套 embedding 建出来的向量索引，否则分数会失真。
                return []
            if indexed_backend != "":
                self._embedding_backend = indexed_backend
            else:
                self._embedding_backend = query_backend
            return self.vector_store.query(query_embedding=query_embedding, limit=limit, node_ids=scope_node_ids)
        except Exception:
            return []

    def _rrf_score(self, rank: int, k: int = 60) -> float:
        return 1.0 / (k + rank)

    def _vector_score_from_distance(self, distance: float) -> float:
        return max(0.0, 1.0 - float(distance))

    def _fuse_search_results(
        self,
        query: str,
        lexical_hits: list[tuple[float, sqlite3.Row]],
        vector_hits: list[VectorSearchHit],
        limit: int,
        *,
        retrieval_profile: RetrievalProfile,
    ) -> list[RetrievalCandidateDebug]:
        """用 RRF 融合 lexical 和 vector 两个召回列表。"""

        # `merged` 以 chunk_id 为键，表示“同一个片段无论被哪一路命中，最终都汇总到一起打分”。
        merged: dict[str, RetrievalCandidateDebug] = {}

        for rank, (lexical_score, row) in enumerate(lexical_hits, start=1):
            candidate = self._row_to_candidate(row)
            merged[candidate.chunk_id] = candidate.model_copy(
                update={
                    "lexical_score": float(lexical_score),
                    "fused_score": self._rrf_score(rank),
                }
            )

        for rank, hit in enumerate(vector_hits, start=1):
            candidate = self._vector_hit_to_candidate(hit)
            vector_score = self._vector_score_from_distance(hit.distance)
            existing = merged.get(candidate.chunk_id)
            if existing is None:
                merged[candidate.chunk_id] = candidate.model_copy(
                    update={
                        "vector_score": vector_score,
                        "fused_score": self._rrf_score(rank),
                    }
                )
                continue
            merged[candidate.chunk_id] = existing.model_copy(
                update={
                    "vector_score": max(existing.vector_score, vector_score),
                    "fused_score": existing.fused_score + self._rrf_score(rank) + 0.02,
                }
            )

        boosted: list[RetrievalCandidateDebug] = []
        for candidate in merged.values():
            boost = self._profile_priority_boost(
                retrieval_profile=retrieval_profile,
                query=query,
                candidate=candidate,
            )
            boosted.append(candidate.model_copy(update={"fused_score": candidate.fused_score + boost}))

        ranked = sorted(
            boosted,
            key=lambda item: (
                item.fused_score,
                item.lexical_score + item.vector_score,
                item.metadata.get("source") == "approved_case",
            ),
            reverse=True,
        )
        return ranked[:limit]

    @property
    def backend_name(self) -> str:
        return f"{self._retrieval_backend}:{self._embedding_backend}"

    @property
    def indexed_embedding_backend(self) -> str:
        return self._load_indexed_embedding_backend() or self._embedding_backend
